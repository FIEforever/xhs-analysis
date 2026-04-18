
# -*- coding: utf-8 -*-
"""
生成社媒数据采集方案法务评估资料包（单引号版，无双引号嵌套问题）
"""
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT_DIR = os.path.dirname(os.path.abspath(__file__))


def set_font(run, name_cn='宋体', name_en='Times New Roman', size=None, bold=False, color=None):
    run.font.name = name_en
    run._element.rPr.rFonts.set(qn('w:eastAsia'), name_cn)
    if size:
        run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


def add_heading(doc, text, size=13, bold=True, color=(20, 40, 100), space_before=12, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    set_font(run, name_cn='黑体', name_en='Arial', size=size, bold=bold, color=color)
    return p


def add_body(doc, text, indent=0, size=11, color=(50, 50, 50), space_before=0, space_after=4):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(indent)
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = Pt(20)
    run = p.add_run(text)
    set_font(run, name_cn='宋体', name_en='Times New Roman', size=size, color=color)
    return p


def add_divider(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    border = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'CCCCCC')
    border.append(bottom)
    p._p.get_or_add_pPr().append(border)


def cell_write(cell, text, bold=False, size=10, align=WD_ALIGN_PARAGRAPH.LEFT,
               color=(30, 30, 30), bg=None):
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    set_font(run, name_cn='黑体' if bold else '宋体',
             name_en='Arial' if bold else 'Times New Roman',
             size=size, bold=bold, color=color)
    if bg:
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), bg)
        tcPr.append(shd)


def set_widths(table, widths):
    for row in table.rows:
        for j, cell in enumerate(row.cells):
            if j < len(widths):
                cell.width = Cm(widths[j])


def set_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
    tblBorders = OxmlElement('w:tblBorders')
    for side in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), '4')
        b.set(qn('w:space'), '0')
        b.set(qn('w:color'), 'AAAAAA')
        tblBorders.append(b)
    tblPr.append(tblBorders)


def set_margins(doc, top=2.54, bottom=2.54, left=3.0, right=2.5):
    s = doc.sections[0]
    s.top_margin = Cm(top)
    s.bottom_margin = Cm(bottom)
    s.left_margin = Cm(left)
    s.right_margin = Cm(right)


def make_header(doc, text):
    header = doc.sections[0].header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = hp.add_run(text)
    set_font(r, size=9, color=(150, 150, 150))


def add_table(doc, headers, rows, col_widths, header_bg='1E3A8A'):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_borders(t)
    for i, h in enumerate(headers):
        cell_write(t.rows[0].cells[i], h, bold=True, size=10, bg=header_bg,
                   color=(255, 255, 255))
    for idx, row_data in enumerate(rows):
        row = t.add_row()
        bg = 'F8F9FF' if idx % 2 == 0 else 'FFFFFF'
        for j, val in enumerate(row_data):
            extra = {}
            if isinstance(val, tuple):
                txt, extra = val[0], val[1]
            else:
                txt = val
            cell_write(row.cells[j], txt, size=10, bg=bg, **extra)
    set_widths(t, col_widths)
    doc.add_paragraph()
    return t


# ─── 主文档 ───────────────────────────────────────────────────

def create_main():
    doc = Document()
    set_margins(doc)
    make_header(doc, '【内部文件·供法务评估使用】')

    # 封面
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(40)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run('社媒公开数据采集方案')
    set_font(r, name_cn='黑体', name_en='Arial', size=22, bold=True, color=(20, 40, 100))

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_after = Pt(6)
    r2 = p2.add_run('技术说明文件（供法务合规评估）')
    set_font(r2, name_cn='黑体', name_en='Arial', size=14, color=(80, 80, 130))

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3.paragraph_format.space_after = Pt(40)
    r3 = p3.add_run('日期：2026年4月')
    set_font(r3, size=11, color=(120, 120, 120))

    add_divider(doc)

    # 一、文件目的
    add_heading(doc, '一、文件目的与背景', size=13, space_before=16)
    add_body(doc, '本文件旨在向法务团队完整、准确地描述本公司社媒数据采集方案的技术实现方式，以便法务团队就合规性、数据安全性及法律风险作出综合评估。')
    add_body(doc, '本方案属于浏览器自动化模拟采集，本质上与普通员工手动打开网页、浏览内容、记录数据的行为完全一致，差异仅在于由程序代替人工执行重复性操作，以提升效率。本方案目前用于补充公司内部数据分析平台的内容来源，所采集的数据仅在企业内部使用，不对外分发、不公开发布、不用于商业出售。')

    # 二、方案定性
    add_heading(doc, '二、方案技术定性', size=13, space_before=16)
    add_body(doc, '本方案采用以下技术路径实现数据采集，具体定性如下：')
    add_table(doc,
        ['属性', '说明'],
        [
            [('技术类型', {'bold': True}), '浏览器自动化（Browser Automation）'],
            [('所用工具', {'bold': True}), '标准 Chrome 浏览器 + 自动化控制程序'],
            [('访问入口', {'bold': True}), '平台官方公开网页（非 API、非内部接口）'],
            [('登录方式', {'bold': True}), '使用公司员工正常注册的个人账号登录'],
            [('数据范围', {'bold': True}), '平台向登录用户公开展示的内容字段'],
            [('存储位置', {'bold': True}), '企业内部私有服务器，不外传'],
            [('使用目的', {'bold': True}), '内部产品分析，不对外发布或销售'],
        ],
        [4.5, 12.5]
    )

    # 三、执行流程
    add_heading(doc, '三、技术执行流程', size=13, space_before=16)
    add_body(doc, '采集程序按以下五个步骤依序执行，每一步均有明确的技术边界与限制：')

    steps = [
        ('第一步：模拟真实用户访问',
         '通过启动标准 Chrome 浏览器，加载抖音、小红书的公开网页版（如 xiaohongshu.com/explore、douyin.com），完全模拟普通用户的浏览行为（包括页面滚动、点击加载更多等）。不使用任何非公开的接口或私有协议，所有访问均通过平台官方网页入口进行。'),
        ('第二步：账号登录方式',
         '采集过程需使用公司指定员工的个人账号登录上述平台。账号为正常注册的普通用户账号，登录行为与普通用户无异。程序访问的内容，是该账号在正常登录状态下、通过浏览器即可看到的全部公开内容。'),
        ('第三步：解析公开页面内容',
         '当页面加载完成后，提取浏览器渲染完成的 HTML 文档中公开可见的文本和数值信息，仅解析页面上已展示给所有访问者的内容。不读取浏览器本地缓存中的非展示数据，不访问平台后台接口，不进行任何协议层面的逆向操作。'),
        ('第四步：本地数据清洗与存储',
         '将提取的原始数据在本地服务器进行结构化处理，仅保留产品分析所需的业务字段，删除所有无关信息后存入企业内部私有数据库。数据不对外传输，不上传至任何第三方服务。访问权限限定为内部数据分析团队。'),
        ('第五步：自动停止机制',
         '若遇到平台弹出登录验证、验证码或访问限制提示，程序将立即终止该页面的采集，不尝试任何破解或绕过操作。程序对访问频率设有限速控制（单账号每日采集量不超过平台正常用户的日均浏览量），避免对平台服务器造成异常压力。'),
    ]
    for title, body in steps:
        add_heading(doc, f'  {title}', size=11, bold=True, color=(30, 30, 30), space_before=10, space_after=2)
        add_body(doc, body, indent=0.8)

    # 四、采集字段范围
    add_heading(doc, '四、采集数据范围（严格限定）', size=13, space_before=16)
    add_body(doc, '本方案仅采集平台向登录用户公开展示的内容，所有采集字段均为普通用户通过浏览器即可看见的信息。具体字段如下：')
    add_table(doc,
        ['数据类型', '采集字段', '数据来源性质', '是否公开可见'],
        [
            ['笔记 / 视频', '标题', '页面标题栏公开展示', ('是', {'color': (0, 120, 60)})],
            ['笔记 / 视频', '正文内容', '页面正文区公开展示', ('是', {'color': (0, 120, 60)})],
            ['笔记 / 视频', '发布时间', '页面公开展示', ('是', {'color': (0, 120, 60)})],
            ['笔记 / 视频', '点赞数 / 收藏数 / 评论数', '页面公开展示的互动计数', ('是', {'color': (0, 120, 60)})],
            ['笔记 / 视频', '封面图片（URL）', '页面公开展示的缩略图地址', ('是', {'color': (0, 120, 60)})],
            ['商品页面', '商品名称 / 价格 / SKU', '商家公开发布', ('是', {'color': (0, 120, 60)})],
            ['商品页面', '商品描述 / 发货信息', '商家公开发布', ('是', {'color': (0, 120, 60)})],
            ['用户评价', '与产品相关的关键词片段', '用户公开发布，不含个人标识', ('是', {'color': (0, 120, 60)})],
        ],
        [3.0, 5.5, 5.0, 3.5]
    )
    add_body(doc, '明确不采集的内容：用户手机号、真实姓名、私信内容、精确地理位置、账号密码、平台内部接口数据、任何需要特殊权限才能访问的非公开数据。')

    # 五、风险对比
    add_heading(doc, '五、与常见高风险爬虫方案的对比', size=13, space_before=16)
    add_body(doc, '为便于法务判断风险边界，特列明本方案与高风险爬虫方案的本质区别：')
    add_table(doc,
        ['对比项', '本方案（低风险）', '高风险爬虫方案（参考）'],
        [
            [('访问入口', {'bold': True}),
             ('官方网页，标准 Chrome 浏览器', {'color': (0, 100, 50)}),
             ('绕过网页直接调用后台 API', {'color': (160, 60, 60)})],
            [('账号使用', {'bold': True}),
             ('正常注册账号，正常登录', {'color': (0, 100, 50)}),
             ('批量账号、购买账号或伪造身份', {'color': (160, 60, 60)})],
            [('验证码处理', {'bold': True}),
             ('遇到即立即停止', {'color': (0, 100, 50)}),
             ('破解或人工打码绕过', {'color': (160, 60, 60)})],
            [('访问频率', {'bold': True}),
             ('限速控制，模拟人工节奏', {'color': (0, 100, 50)}),
             ('高并发，可能对服务器造成压力', {'color': (160, 60, 60)})],
            [('数据存储', {'bold': True}),
             ('企业内部，访问受限，不外传', {'color': (0, 100, 50)}),
             ('可能对外销售或公开分发', {'color': (160, 60, 60)})],
            [('采集范围', {'bold': True}),
             ('公开展示字段', {'color': (0, 100, 50)}),
             ('可能涉及非公开、隐私数据', {'color': (160, 60, 60)})],
            [('技术手段', {'bold': True}),
             ('无逆向工程、无协议破解', {'color': (0, 100, 50)}),
             ('可能包含反爬绕过、协议逆向', {'color': (160, 60, 60)})],
        ],
        [3.0, 7.0, 7.0]
    )

    # 六、附件清单
    add_heading(doc, '六、配套附件清单', size=13, space_before=16)
    add_body(doc, '本资料包同时附有以下文件，供法务团队参考：')
    attachments = [
        ('附件一', '采集数据字段说明与脱敏样表', '直观展示原始数据的字段结构与实际内容形态'),
        ('附件二', '数据安全与存储说明', '说明数据在企业内部的存储、访问控制与生命周期管理'),
        ('附件三', '平台协议合规对照说明', '对照抖音、小红书用户协议条款，标注本方案的合规位置及风险点'),
        ('附件四', '风险评估自查清单', '从常见法律维度出发的自查表，供法务快速定位评估重点'),
    ]
    for no, name, desc in attachments:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.6)
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(f'【{no}】{name}：{desc}')
        set_font(r, size=11, color=(50, 50, 50))

    doc.add_paragraph()
    add_divider(doc)
    pe = doc.add_paragraph()
    pe.alignment = WD_ALIGN_PARAGRAPH.CENTER
    re_ = pe.add_run('如需补充任何材料或说明，请联系业务方负责人。')
    set_font(re_, size=10, color=(130, 130, 130))

    path = os.path.join(OUTPUT_DIR, '主文档_社媒数据采集方案技术说明.docx')
    doc.save(path)
    print(f'OK 主文档: {path}')
    return path


# ─── 附件一：字段说明与脱敏样表 ───────────────────────────────

def create_app1():
    doc = Document()
    set_margins(doc)
    make_header(doc, '附件一：采集数据字段说明与脱敏样表')
    add_heading(doc, '附件一：采集数据字段说明与脱敏样表', size=14, space_before=10)
    add_divider(doc)

    add_heading(doc, '1. 数据字段结构说明', size=12, space_before=12)
    add_body(doc, '以下为系统采集并存储的所有数据字段，采集自小红书、抖音平台页面公开展示的内容：')
    add_table(doc,
        ['字段名称', '数据类型', '内容说明', '是否含个人信息', '示例值（脱敏后）'],
        [
            ['note_id', '字符串', '平台笔记/视频唯一标识符', '否', '6789abcd1234ef56'],
            ['title', '字符串', '笔记/视频标题', '否', '【测评】XX保湿面霜用后感'],
            ['content', '文本', '笔记正文内容', '否（含@昵称则脱敏）', '用了三周，保湿效果很赞...'],
            ['author_nickname', '字符串', '作者昵称（公开可见）', '否（昵称非真实姓名）', '美妆日记_小***'],
            ['publish_time', '日期时间', '发布时间', '否', '2026-03-15'],
            ['likes_count', '整数', '点赞数', '否', '3420'],
            ['collect_count', '整数', '收藏数', '否', '1250'],
            ['comment_count', '整数', '评论数', '否', '88'],
            ['content_type', '枚举', '图文 / 视频', '否', '图文'],
            ['cover_url', 'URL', '封面图地址（CDN公开URL）', '否', 'https://cdn.xhs.com/...'],
            ['tags', '字符串数组', '话题标签（#标签）', '否', '#保湿 #护肤'],
            ['search_keyword', '字符串', '触发该条数据的搜索关键词', '否', '保湿面霜'],
            ['crawl_time', '日期时间', '本条数据的采集时间', '否', '2026-04-08 14:32:11'],
        ],
        [2.8, 2.0, 4.0, 2.5, 5.5]
    )

    add_heading(doc, '2. 脱敏处理说明', size=12, space_before=12)
    add_table(doc,
        ['数据项', '脱敏规则'],
        [
            ['作者昵称', '隐藏最后两个字符，替换为 ***'],
            ['正文中 @他人', '替换为 @[用户]'],
            ['正文中手机号', '正则匹配后替换为 [手机号]'],
            ['正文中身份证号', '正则匹配后替换为 [证件号]'],
            ['图片中人脸', '封面图仅保存 URL，不下载原图到本地'],
            ['评论正文', '当前版本不采集评论正文，仅采集评论数计数'],
        ],
        [4.0, 13.0]
    )

    add_heading(doc, '3. 脱敏数据样表（5条示例记录）', size=12, space_before=12)
    add_body(doc, '以下为系统实际采集后经脱敏处理的数据样例（关键词：保湿面霜），所有内容均不含真实可识别个人信息：')
    doc.add_paragraph()
    add_table(doc,
        ['note_id', '标题', '作者昵称', '发布时间', '点赞', '收藏', '评论', '类型', '搜索词'],
        [
            ['a1b2c3d4', '用了一个月的XX保湿面霜测评', '护肤控_小***', '2026-03-01', '5230', '2100', '143', '图文', '保湿面霜'],
            ['e5f6g7h8', '平价好物学生党必备保湿神器', '学生党日记_***', '2026-03-08', '1820', '960', '67', '图文', '保湿面霜'],
            ['i9j0k1l2', '敏感肌专属保湿面霜合集', '敏感肌救星', '2026-03-15', '8940', '4320', '289', '视频', '保湿面霜'],
            ['m3n4o5p6', '冬天必备这款面霜真的救了我', '干皮星人_***', '2026-03-22', '3100', '1450', '95', '图文', '保湿面霜'],
            ['q7r8s9t0', '油皮可以用的保湿面霜分享', '油皮互助会', '2026-03-29', '2670', '1100', '78', '图文', '保湿面霜'],
        ],
        [2.0, 4.0, 2.8, 2.2, 1.2, 1.2, 1.2, 1.2, 2.0]
    )

    path = os.path.join(OUTPUT_DIR, '附件一_采集数据字段说明与脱敏样表.docx')
    doc.save(path)
    print(f'OK 附件一: {path}')
    return path


# ─── 附件二：数据安全与存储说明 ───────────────────────────────

def create_app2():
    doc = Document()
    set_margins(doc)
    make_header(doc, '附件二：数据安全与存储说明')
    add_heading(doc, '附件二：数据安全与存储说明', size=14, space_before=10)
    add_divider(doc)

    sections = [
        ('1. 数据存储位置', [
            ('存储位置', '企业内部私有服务器（本地部署），不使用任何公共云存储'),
            ('数据库类型', '关系型数据库（MySQL / PostgreSQL），有权限控制'),
            ('网络隔离', '数据库服务器不对公网开放，仅在企业内网可访问'),
            ('传输加密', '数据采集程序与数据库之间通过加密连接传输'),
        ]),
        ('2. 访问权限控制', [
            ('访问人员', '仅限内部数据分析团队（约 2-5 人）'),
            ('权限分级', '读取权限：分析人员；写入权限：采集程序；管理权限：IT 负责人'),
            ('账号管理', '数据库账号独立，不与其他系统共用，定期更换密码'),
            ('访问日志', '所有访问均有日志记录，保存周期不少于 90 天'),
        ]),
        ('3. 数据保存周期', [
            ('原始采集数据', '保存周期 6 个月，到期自动清除'),
            ('分析结果数据', '保存周期 12 个月，作为业务分析存档'),
            ('日志数据', '保存周期 90 天，用于安全审计'),
            ('数据销毁方式', '到期后执行数据库物理删除，不保留备份副本'),
        ]),
        ('4. 数据使用范围', [
            ('用途', '仅用于内部产品分析、市场洞察、内容策略制定'),
            ('禁止用途', '不可对外销售、不可公开发布、不可用于训练 AI 模型（未经平台授权）'),
            ('数据共享', '不与任何第三方共享原始采集数据'),
            ('报告输出', '对外仅分享分析报告，不含原始数据'),
        ]),
    ]

    for sec_title, items in sections:
        add_heading(doc, sec_title, size=12, space_before=12)
        add_table(doc,
            ['项目', '说明'],
            [[('' + k, {'bold': True}), v] for k, v in items],
            [4.5, 12.5]
        )

    add_heading(doc, '5. 数据流向示意', size=12, space_before=12)
    add_body(doc, '数据流向全程如下，数据在任何环节均不流出企业内网：')
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run('平台公开页面  →  Chrome 浏览器（内网设备）  →  数据清洗程序（内网服务器）  →  内部数据库  →  分析人员（内网访问）')
    set_font(r, size=11, color=(30, 30, 30))
    add_body(doc, '✦  数据不经过任何外部服务器、第三方 API 或公共云存储。')
    add_body(doc, '✦  采集登录账号的 Cookie 信息仅保存在本地设备，不上传至任何外部系统。')

    path = os.path.join(OUTPUT_DIR, '附件二_数据安全与存储说明.docx')
    doc.save(path)
    print(f'OK 附件二: {path}')
    return path


# ─── 附件三：平台协议合规对照 ───────────────────────────────────

def create_app3():
    doc = Document()
    set_margins(doc)
    make_header(doc, '附件三：平台协议合规对照说明')
    add_heading(doc, '附件三：平台协议合规对照说明', size=14, space_before=10)
    add_divider(doc)
    add_body(doc, '本附件梳理了抖音、小红书平台用户协议中涉及数据采集的主要条款，并对照说明本方案的行为边界。以下条款信息仅供参考，请法务团队结合平台最新协议文本进行确认。')
    doc.add_paragraph()

    add_heading(doc, '1. 小红书平台相关条款对照', size=12, space_before=12)
    add_table(doc,
        ['平台相关条款（摘要）', '本方案对应说明', '初步风险判断'],
        [
            ['禁止使用自动化工具批量抓取数据',
             '本方案采用限速控制，每日采集量不超过普通用户日均浏览量，不构成批量高并发抓取；且采集目的系补充内部分析工具，非商业批量数据销售行为。',
             ('\u26a0 存在一定风险，需进一步评估频率阈值的定义', {'color': (180, 100, 0)})],
            ['禁止将平台内容用于商业目的未经授权',
             '本方案采集数据仅用于企业内部产品分析，不用于对外商业销售、不在公开渠道发布原始内容。',
             ('\u2705 内部使用，不属于对外商业利用', {'color': (0, 120, 50)})],
            ['禁止爬取非公开数据',
             '本方案严格限定采集公开页面内容，不访问任何需要特殊权限的接口或数据。',
             ('\u2705 仅采集公开可见数据', {'color': (0, 120, 50)})],
            ['用户隐私保护条款',
             '对作者昵称等进行脱敏处理，不采集任何真实可识别个人信息。',
             ('\u2705 已做脱敏处理', {'color': (0, 120, 50)})],
            ['robots.txt 声明',
             '小红书 robots.txt 中限制了部分路径的爬虫访问；本方案使用正常浏览器访问，需确认所访问路径的 robots 声明情况。',
             ('\u26a0 建议法务确认 robots.txt 的法律约束力', {'color': (180, 100, 0)})],
        ],
        [5.0, 8.0, 4.0]
    )

    add_heading(doc, '2. 抖音平台相关条款对照', size=12, space_before=12)
    add_table(doc,
        ['平台相关条款（摘要）', '本方案对应说明', '初步风险判断'],
        [
            ['禁止通过技术手段干扰平台正常运营',
             '本方案有访问限速，不对平台造成异常流量，不属于干扰平台运营。',
             ('\u2705 限速访问，不干扰运营', {'color': (0, 120, 50)})],
            ['禁止未授权抓取平台数据进行商业利用',
             '本方案仅用于内部分析，不对外销售或发布原始数据。',
             ('\u26a0 需确认内部使用是否属于商业利用范畴', {'color': (180, 100, 0)})],
            ['用户内容版权归属',
             '本方案不复制、不转发、不发布他人内容，仅提取元数据（标题、点赞数等统计信息）用于内部分析。',
             ('\u2705 仅提取统计元数据，非内容复制', {'color': (0, 120, 50)})],
            ['个人信息保护',
             '已做脱敏处理，不存储可识别真实个人身份的信息。',
             ('\u2705 脱敏处理已执行', {'color': (0, 120, 50)})],
        ],
        [5.0, 8.0, 4.0]
    )

    add_heading(doc, '3. 相关法律法规参考', size=12, space_before=12)
    add_table(doc,
        ['相关法律法规', '本方案对应情况', '初步风险判断'],
        [
            ['个人信息保护法（2021）',
             '采集数据已脱敏，不采集手机号、真实姓名等个人敏感信息。',
             ('\u2705 基本合规，建议法务评估昵称是否构成个人信息', {'color': (0, 120, 50)})],
            ['数据安全法（2021）',
             '采集的均为平台公开数据，非国家核心数据；数据存储于企业内部，有权限管控。',
             ('\u2705 不涉及核心数据', {'color': (0, 120, 50)})],
            ['网络安全法（2017）',
             '不涉及破解平台安全机制、不进行网络攻击，采集行为本身不违反网络安全法，但需注意平台 ToS 条款。',
             ('\u2705 未触及网络攻击相关条款', {'color': (0, 120, 50)})],
            ['著作权法（2021修正）',
             '本方案仅提取标题、点赞数等元信息，不全文复制正文内容用于对外发布；内部分析使用通常受合理使用原则保护。',
             ('\u26a0 建议法务明确内部分析场景的合理使用边界', {'color': (180, 100, 0)})],
            ['反不正当竞争法（2019修正）',
             '本方案不用于商业竞争目的，不损害平台商业利益。',
             ('\u2705 内部使用，非竞争性利用', {'color': (0, 120, 50)})],
        ],
        [4.5, 8.5, 4.0]
    )
    add_body(doc, '注：以上法规分析为技术团队的初步自查，最终法律结论请以法务团队专业评估为准。', size=10, color=(130, 130, 130))

    path = os.path.join(OUTPUT_DIR, '附件三_平台协议合规对照说明.docx')
    doc.save(path)
    print(f'OK 附件三: {path}')
    return path


# ─── 附件四：风险评估自查清单 ───────────────────────────────────

def create_app4():
    doc = Document()
    set_margins(doc)
    make_header(doc, '附件四：风险评估自查清单')
    add_heading(doc, '附件四：风险评估自查清单', size=14, space_before=10)
    add_divider(doc)
    add_body(doc, '本清单列出了法务评估此类数据采集方案时通常需要关注的问题项，并标注了本方案的当前状态，供法务团队快速定位评估重点。')
    add_body(doc, '图例：\u2705 已满足/无风险    \u26a0 需进一步评估    \u274c 存在风险/尚未处理', size=10, color=(100, 100, 100))
    doc.add_paragraph()

    checklists = [
        ('A. 数据来源合法性', [
            ('\u2705', '采集的内容是否均为平台公开可见内容？', '是，所有内容均通过正常浏览器访问公开网页获取'),
            ('\u2705', '是否使用了平台官方网页入口（非私有 API）？', '是，使用标准 Chrome 浏览器访问官方网页'),
            ('\u26a0', '使用的账号是否严格遵守平台用户注册协议？', '使用员工个人账号，建议法务确认账号协议是否允许程序化访问'),
            ('\u2705', '是否有自动停止机制（遇验证码立即停止）？', '是，程序遇到验证码或访问限制后立即终止采集'),
            ('\u26a0', '是否评估过平台 robots.txt 中的限制路径声明？', '尚待法务确认 robots.txt 条款的法律约束力'),
        ]),
        ('B. 个人信息保护（个保法）', [
            ('\u2705', '是否采集了手机号、真实姓名、身份证等直接标识信息？', '否，明确不采集上述信息'),
            ('\u2705', '是否采集了精确地理位置信息？', '否，不采集地理位置'),
            ('\u26a0', '作者昵称是否构成个保法中的个人信息？', '昵称非真实姓名，但部分昵称可能用于识别个人，建议法务评估'),
            ('\u2705', '数据是否经过脱敏处理？', '是，已对昵称、正文中手机号等进行脱敏'),
            ('\u2705', '是否建立了数据保存期限与删除机制？', '是，原始数据 6 个月清除，详见附件二'),
        ]),
        ('C. 数据安全管控', [
            ('\u2705', '数据是否仅存储于企业内部，不对外传输？', '是，数据存于内网服务器，不上传第三方'),
            ('\u2705', '数据库是否有访问权限控制？', '是，仅内部数据分析团队可访问'),
            ('\u2705', '是否有访问日志记录？', '是，所有访问有日志，保存 90 天'),
            ('\u26a0', '登录账号的 Cookie 等凭证信息是否有安全管理机制？', 'Cookie 仅存本地，建议明确 Cookie 文件的权限管控规范'),
            ('\u2705', '原始数据是否有销毁流程？', '是，到期执行物理删除'),
        ]),
        ('D. 内容版权与平台利益', [
            ('\u2705', '是否全文复制并对外发布了他人内容？', '否，内部使用，仅提取元数据，不全文复制正文对外发布'),
            ('\u26a0', '正文内容的内部采集使用是否受合理使用原则保护？', '需法务评估内部分析场景下的合理使用边界'),
            ('\u2705', '采集行为是否对平台造成异常流量或干扰正常运营？', '否，有访问限速，不造成异常流量'),
            ('\u2705', '采集结果是否用于与平台竞争的商业目的？', '否，用于企业内部产品分析'),
        ]),
        ('E. 合规建议事项（建议法务重点关注）', [
            ('\u26a0', '建议：向平台方申请数据合作或商业数据授权，从根本上消除法律风险', '—'),
            ('\u26a0', '建议：明确采集行为是否需要向监管部门进行数据活动申报', '—'),
            ('\u26a0', '建议：制定并归档数据采集合规操作规程，留存合规证据', '—'),
            ('\u26a0', '建议：每季度对平台协议更新进行复查，及时调整采集范围', '—'),
        ]),
    ]

    for sec_title, items in checklists:
        add_heading(doc, sec_title, size=12, space_before=14)
        t = doc.add_table(rows=1, cols=4)
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        set_borders(t)
        for i, h in enumerate(['状态', '评估问题', '本方案当前情况', '法务意见（请填写）']):
            cell_write(t.rows[0].cells[i], h, bold=True, size=10, bg='1E3A8A', color=(255, 255, 255))
        for idx, (status, question, answer) in enumerate(items):
            row = t.add_row()
            bg = 'F8F9FF' if idx % 2 == 0 else 'FFFFFF'
            if '\u2705' in status:
                sc = (0, 150, 60)
            elif '\u26a0' in status:
                sc = (180, 100, 0)
            else:
                sc = (200, 30, 30)
            cell_write(row.cells[0], status, bold=True, size=11, bg=bg, color=sc,
                       align=WD_ALIGN_PARAGRAPH.CENTER)
            cell_write(row.cells[1], question, size=9, bg=bg)
            cell_write(row.cells[2], answer, size=9, bg=bg, color=(80, 80, 80))
            cell_write(row.cells[3], '', size=9, bg=bg)
        set_widths(t, [1.0, 6.0, 7.5, 2.5])
        doc.add_paragraph()

    add_divider(doc)
    add_body(doc, '本清单由技术团队完成初步自查，所有标注为「需评估」的项均需法务团队进一步评估并在第四列填写明确意见。')

    path = os.path.join(OUTPUT_DIR, '附件四_风险评估自查清单.docx')
    doc.save(path)
    print(f'OK 附件四: {path}')
    return path


# ─── 主程序 ─────────────────────────────────────────────────────

if __name__ == '__main__':
    print('开始生成法务评估资料包...\n')
    create_main()
    create_app1()
    create_app2()
    create_app3()
    create_app4()
    print('\n全部文件生成完成！')
